"""
Автоматическая проверка заданий студентов с использованием Gemini API.
Адаптированная версия для проверки Word документов и HTML текстов.

Особенности:
- Использование официального Gemini API
- Ротация 6 API ключей для избежания квот
- Автоматические задержки для соблюдения лимитов
- Обработка ошибок и повторные попытки
- Бинарная оценка: зачтено/не зачтено
- Сохранение результатов в result.txt и итоговую ведомость
"""

import asyncio
import json
import logging
import os
import threading
from threading import Thread
from typing import Callable, List
import time
import random
import re

import aiofiles
import pandas as pd
from bs4 import BeautifulSoup
from colorama import init
from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type
from tqdm import tqdm
from urllib.parse import urlparse

# Импорт Gemini API
from google import genai
from google.genai import types

# Инициализация colorama для цветного вывода в терминале
init(autoreset=True)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# API ключи Gemini для ротации
API_KEYS = [
    "AIzaSyBZERjXBrLTMniOGnajrIrzoZvBkcx8_dE",
    "AIzaSyBVyS1pW5X7WW6qKLbd4wZxGYy4oLV5YGA",
    "AIzaSyDNRbu7_VJmp1wwY_9Ziz8VqXRQo7CNYJo",
    "AIzaSyB3JtO6koAQZBsu7_S3jgryn81asEHSQqQ",
    "AIzaSyBS_L7_FwnFhK5M3GglsJc1jaUVS4Uv_T0",
    "AIzaSyBj_YDJXk3QnhIH8FdIZkkPOnxNfirSyWY",
    "AIzaSyBP-sSpfUVTAWXd-prxEvdbBeLCaw_NCsk",
    "AIzaSyC6exCJDufHqzSj10vMGAJ0GgIT4WnmWQA", 
    "AIzaSyATbFpgBVoYmxr0oaaRI9tEstkXM5aHmwY",
    "AIzaSyDNyyAGQVSoqw1gDov03iRNtx9-9vH28iI",
    "AIzaSyDDzd5N3Z_be3OMhsPQGhF4YHc9bM5RAMo",
    "AIzaSyAWQjDRCIwGfn74jfaNaZjSUBsvAu9hpKw"
]

# Список моделей Gemini для перебора
AVAILABLE_MODELS = [
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite"
]

# Настройки задержек для соблюдения квот
MIN_DELAY = 10.0  # Минимальная задержка между запросами (секунды)
MAX_DELAY = 20.0  # Максимальная задержка между запросами (секунды)
REQUEST_TIMEOUT = 120  # Таймаут запроса (секунды)


class GeminiClient:
    """Клиент для работы с Gemini API с автоматическими задержками и ротацией ключей."""

    def __init__(self, api_key: str = None):
        self.api_keys = API_KEYS.copy()
        self.current_key_index = 0
        self.api_key = api_key or self.api_keys[self.current_key_index]
        self._http_options = self._build_http_options()
        self.client = genai.Client(api_key=self.api_key, http_options=self._http_options)
        self.last_request_time = 0
        self.key_usage_count = {key: 0 for key in self.api_keys}

    @staticmethod
    def _build_http_options():
        """Собирает настройки HTTP с учётом прокси из окружения."""

        proxies = {}

        raw_proxy = (
            os.environ.get("HTTPS_PROXY")
            or os.environ.get("https_proxy")
            or os.environ.get("HTTP_PROXY")
            or os.environ.get("http_proxy")
        )

        if not raw_proxy:
            host = os.environ.get("PROXY_HOST")
            port = os.environ.get("PROXY_PORT")
            if host and port:
                user = os.environ.get("PROXY_USER")
                password = os.environ.get("PROXY_PASS")
                credentials = f"{user}:{password}@" if user and password else ""
                raw_proxy = f"socks5h://{credentials}{host}:{port}"

        if raw_proxy:
            parsed = urlparse(raw_proxy)
            scheme = parsed.scheme.lower()
            if not scheme:
                raw_proxy = f"socks5h://{raw_proxy}"
            elif scheme in {"http", "https"}:
                raw_proxy = raw_proxy.replace(f"{parsed.scheme}://", "socks5h://", 1)

            proxies = {"http": raw_proxy, "https": raw_proxy}
            logger.info("Используется SOCKS5 прокси для Gemini API: %s", raw_proxy)

        if not proxies:
            return None

        client_args = {
            "proxies": proxies,
            "http2": False,
        }

        return types.HttpOptions(
            client_args=client_args,
            async_client_args=client_args,
        )
        
    async def _wait_for_rate_limit(self):
        """Ожидание для соблюдения лимитов API."""
        current_time = time.time()
        time_since_last_request = current_time - self.last_request_time
        
        if time_since_last_request < MIN_DELAY:
            wait_time = MIN_DELAY - time_since_last_request + random.uniform(0, MAX_DELAY - MIN_DELAY)
            logger.debug(f"Ожидание {wait_time:.2f} секунд для соблюдения лимитов API...")
            await asyncio.sleep(wait_time)
        
        self.last_request_time = time.time()
    
    def _rotate_api_key(self):
        """Переключение на следующий API ключ."""
        self.current_key_index = (self.current_key_index + 1) % len(self.api_keys)
        self.api_key = self.api_keys[self.current_key_index]
        self.client = genai.Client(api_key=self.api_key, http_options=self._http_options)
        logger.info(f"Переключение на API ключ {self.current_key_index + 1}/{len(self.api_keys)}")
    
    async def _handle_quota_error(self, error_msg: str, attempt: int) -> int:
        """Обработка ошибок квот с экспоненциальной задержкой и ротацией ключей."""
        if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
            # Попробуем переключить ключ
            self._rotate_api_key()
            
            # Извлекаем время ожидания из сообщения об ошибке
            import re
            retry_match = re.search(r'retry in (\d+(?:\.\d+)?)s', error_msg)
            if retry_match:
                wait_time = float(retry_match.group(1))
                # Добавляем дополнительное время для стабильности
                wait_time += 10
            else:
                # Экспоненциальная задержка: 60s, 120s, 240s...
                wait_time = 60 * (2 ** (attempt - 1))
            
            logger.warning(f"🔄 Превышена квота API. Переключение на ключ {self.current_key_index + 1}/{len(self.api_keys)} и ожидание {wait_time:.1f} секунд...")
            await asyncio.sleep(wait_time)
            return attempt + 1
        return attempt
    
    async def generate_content(self, text: str, model: str = None) -> str:
        """
        Генерация контента с помощью Gemini API.
        
        Args:
            text: Текст для обработки
            model: Модель для использования (по умолчанию первая из списка)
            
        Returns:
            Сгенерированный текст
        """
        if model is None:
            model = AVAILABLE_MODELS[0]
            
        await self._wait_for_rate_limit()
        
        try:
            logger.debug(f"Отправка запроса к модели {model} с ключом {self.current_key_index + 1}")
            
            contents = [
                types.Content(
                    role="user",
                    parts=[
                        types.Part.from_text(text=text),
                    ],
                ),
            ]
            
            generate_content_config = types.GenerateContentConfig(
                temperature=0.1,  # Низкая температура для более детерминированных ответов
                max_output_tokens=8192,  # Максимальное количество токенов в ответе
            )
            
            # Собираем полный ответ
            response = self.client.models.generate_content(
                model=model,
                contents=contents,
                config=generate_content_config,
            )

            # Собираем текст ответа
            full_response = getattr(response, "text", "") or ""
            if not full_response and getattr(response, "candidates", None):
                parts: list[str] = []
                for candidate in response.candidates:
                    content = getattr(candidate, "content", None)
                    if not content:
                        continue
                    for part in getattr(content, "parts", []):
                        text_part = getattr(part, "text", None)
                        if text_part:
                            parts.append(text_part)
                full_response = "".join(parts)

            # Увеличиваем счетчик использования ключа
            self.key_usage_count[self.api_key] += 1
            logger.debug(
                "Получен ответ от модели %s (ключ %s, использований: %s)",
                model,
                self.current_key_index + 1,
                self.key_usage_count[self.api_key],
            )
            return full_response.strip()
            
        except Exception as e:
            logger.warning(f"Ошибка при запросе к модели {model} с ключом {self.current_key_index + 1}: {e}")
            raise


async def extract_text_from_html(html_path: str) -> str:
    """
    Извлекает текст из HTML файла.
    """
    try:
        async with aiofiles.open(html_path, 'r', encoding='utf-8') as f:
            html_content = await f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        text = soup.get_text(separator=' ', strip=True)
        return text
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из HTML {html_path}: {e}")
        return ""


async def extract_text_from_word(word_path: str) -> str:
    """
    Извлекает текст из Word документа.
    """
    try:
        # Попытка использовать python-docx для качественного извлечения
        try:
            from docx import Document
            doc = Document(word_path)
            text_parts = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            full_text = '\n'.join(text_parts)
            
            # Если текст слишком короткий, возможно это шаблон
            if len(full_text.strip()) < 100:
                logger.warning(f"Возможно, это шаблон или пустой документ: {word_path}")
                return full_text
            
            return full_text
            
        except ImportError:
            logger.warning("python-docx не установлен, используется базовое извлечение")
        except Exception as e:
            logger.warning(f"Ошибка при использовании python-docx: {e}, используется базовое извлечение")
        
        # Базовое извлечение текста как fallback
        async with aiofiles.open(word_path, 'rb') as f:
            content = await f.read()
        
        # Базовое извлечение текста (упрощенная версия)
        text = content.decode('utf-8', errors='ignore')
        # Очистка от служебных символов Word
        text = re.sub(r'[^\w\s\u0400-\u04FF.,!?;:()\-]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из Word {word_path}: {e}")
        return ""


async def extract_text_from_pdf(pdf_path: str) -> str:
    """Извлекает текст из PDF документа.

    Порядок попыток:
    1) PyPDF2 — лёгкий и без внешних бинарников;
    2) Фолбэк: двоичное чтение с грубой очисткой (на случай отсутствия библиотек).
    """
    try:
        try:
            import PyPDF2  # noqa: WPS433
            text_parts: list[str] = []
            with open(pdf_path, 'rb') as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:
                        page_text = ""
                    if page_text.strip():
                        text_parts.append(page_text)
            text = "\n".join(text_parts).strip()
            return text
        except ImportError:
            logger.warning("PyPDF2 не установлен, используется базовое извлечение PDF")

        # Фолбэк: читаем как текст с доп. очисткой (качество ниже)
        async with aiofiles.open(pdf_path, 'rb') as f:
            content = await f.read()
        text = content.decode('utf-8', errors='ignore')
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из PDF {pdf_path}: {e}")
        return ""


async def answer(client: GeminiClient, text: str, prompt: str = '', limit: int = 60,
                 models: list = None) -> str | None:
    """
    Отправляет запрос к модели Gemini и возвращает ответ.
    Перебирает модели из списка, пока не получит валидный ответ.
    Обрабатывает квоты с повторными попытками.
    """
    if models is None:
        models = AVAILABLE_MODELS.copy()

    for model in models:
        attempt = 1
        max_attempts = 10  # Увеличиваем количество попыток для обработки квот
        
        while attempt <= max_attempts:
            try:
                logger.info(f"Попытка {attempt}/{max_attempts} с моделью {model}")
                response = await client.generate_content(f"{prompt}\n{text}", model)
                if response:
                    # Удаляем алфавитно-цифровые символы и лишние пробелы
                    clear_text = re.sub(r"[A-Za-z0-9]", "", response)
                    clear_text = re.sub(r'\s+', ' ', clear_text.strip())
                    if len(clear_text) > limit:
                        logger.info(f"Получен валидный ответ от модели {model}.")
                        return response.strip()
                    else:
                        logger.warning(f"Ответ слишком короткий: {len(clear_text)} символов.")
                        if attempt < max_attempts:
                            await asyncio.sleep(5)
                            attempt += 1
                            continue
            except Exception as e:
                error_msg = str(e)
                logger.warning(f"Ошибка при работе с моделью {model}, попытка {attempt}/{max_attempts}: {e}")
                
                # Обработка ошибок квот
                if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                    attempt = await client._handle_quota_error(error_msg, attempt)
                    if attempt > max_attempts:
                        logger.error(f"Превышено максимальное количество попыток для модели {model}")
                        break
                    continue
                elif "404" in error_msg or "NOT_FOUND" in error_msg:
                    logger.warning(f"Модель {model} недоступна, переход к следующей модели...")
                    break
                else:
                    if attempt < max_attempts:
                        await asyncio.sleep(5)
                        attempt += 1
                        continue
                    else:
                        logger.error(f"Модель {model} завершилась неудачей после {max_attempts} попыток")
                        break

    logger.error("Все модели завершились неудачей.")
    return None


def extract_binary_result(msg: str) -> dict | None:
    """
    Извлекает бинарный результат (зачтено/не зачтено) из ответа AI.
    """
    try:
        logger.debug("Попытка извлечь бинарный результат из ответа...")
        
        # Сначала пытаемся найти JSON структуру
        json_match = re.search(r'\{[^{}]*"result"[^{}]*\}', msg, re.DOTALL)
        if json_match:
            try:
                json_str = json_match.group(0)
                data = json.loads(json_str)
                result = data.get('result', 'не зачтено').lower()
                comment = data.get('comment', 'Комментарий не предоставлен')
                return {
                    "result": result,
                    "comment": comment
                }
            except json.JSONDecodeError:
                logger.warning("Не удалось распарсить JSON, используем fallback")
        
        # Fallback: ищем по ключевым словам
        msg_lower = msg.lower()
        
        # Проверяем "не зачтено" ПЕРЕД "зачтено", чтобы избежать ложных срабатываний
        if "не зачтено" in msg_lower or "незачет" in msg_lower or "не засчитано" in msg_lower:
            result = "не зачтено"
        elif "зачтено" in msg_lower or "зачет" in msg_lower or "засчитано" in msg_lower:
            result = "зачтено"
        else:
            result = "не зачтено"  # По умолчанию
        
        # Извлекаем комментарий
        comment = msg
        if "комментарий:" in msg_lower:
            comment = msg.split("комментарий:")[-1].strip()
        elif "comment:" in msg_lower:
            comment = msg.split("comment:")[-1].strip()
        
        # Если комментарий слишком длинный, обрезаем
        if len(comment) > 500:
            comment = comment[:500] + "..."
        
        return {
            "result": result,
            "comment": comment
        }
        
    except Exception as e:
        logger.error(f"Ошибка при извлечении бинарного результата: {e}")
        return {
            "result": "не зачтено",
            "comment": "Ошибка при обработке ответа"
        }


async def check_ai_generation(client: GeminiClient, text: str, models: list = None) -> dict | None:
    """
    Проверяет работу на признаки AI-генерации.
    """
    if models is None:
        models = AVAILABLE_MODELS.copy()

    logger.info("Проверка на признаки AI-генерации...")

    prompt = f"""
Ты — эксперт по детекции AI-генерации.  
Проанализируй текст на признаки того, что он был создан с помощью искусственного интеллекта.

ТЕКСТ ДЛЯ АНАЛИЗА:
{text}

ОЦЕНИ ТОЛЬКО ОЧЕВИДНЫЕ СЛУЧАИ AI-ГЕНЕРАЦИИ.

ПРИЗНАКИ AI-ГЕНЕРАЦИИ:
1. Текст написан энциклопедическим, академическим или шаблонным стилем, без личных выражений.
2. Используются характерные клише ChatGPT-подобных ответов: «в заключение», «таким образом», «важно отметить», «повышение эффективности» и т.п.
3. В тексте отсутствуют ошибки, личные наблюдения, логические отклонения или естественная вариативность речи.
4. Предложения однотипны по длине и структуре, нет эмоциональной окраски.
5. Текст полностью состоит из обобщений, без конкретных примеров, фактов или личных формулировок.

ПРАВИЛА:
— Если признаки AI-генерации очевидны по нескольким пунктам, укажи ai_detected: true.  
— Если есть сомнения, всегда выбирай ai_detected: false.  
— Не считай AI-генерацией просто грамотный или аккуратный текст — только явно шаблонные случаи.

ФОРМАТ ВЫВОДА:
{{
  "ai_detected": true или false,
  "confidence": "низкая", "средняя" или "высокая",
  "reasons": ["список конкретных признаков"],
  "comment": "детальный анализ стиля и содержания"
}}

Отвечай на русском языке.
"""

    response = await answer(client, text=text, prompt=prompt, models=models)
    if response:
        try:
            # Пытаемся найти JSON в ответе
            json_match = re.search(r'\{[^{}]*"ai_detected"[^{}]*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                data = json.loads(json_str)
                return {
                    "ai_detected": data.get('ai_detected', False),
                    "confidence": data.get('confidence', 'низкая'),
                    "reasons": data.get('reasons', []),
                    "comment": data.get('comment', 'Анализ не выполнен')
                }
        except json.JSONDecodeError:
            logger.warning("Не удалось распарсить JSON ответ для проверки AI")
    
    return None


async def get_binary_evaluation(
    client: GeminiClient,
    text: str,
    template_text: str,
    models: list = None,
    room_prompt: str = "",
    ai_confidence: str | None = None,
    ai_check_enabled: bool = True,
) -> dict | None:
    """
    Запрашивает бинарную оценку у модели.
    """
    if models is None:
        models = AVAILABLE_MODELS.copy()

    logger.info("Запрос бинарной оценки у модели...")

    # Промпт для режима БЕЗ AI проверки - фокус на качестве текста
    if not ai_check_enabled:
        base_prompt = f"""
Ты — эксперт по оценке студенческих работ.  
Оцени работу строго, по существу, без поблажек, но с учётом учебного уровня.  
Результат всегда бинарный: "зачтено" или "не зачтено".  

ВАЖНО: В этой проверке НЕ учитывается фактор AI-генерации. 
Сосредоточься ТОЛЬКО на качестве содержания, логике изложения и оформлении.

КРИТЕРИИ ОЦЕНКИ ИЗ ШАБЛОНА:
{template_text}

РАБОТА СТУДЕНТА:
{text}

---

КРИТЕРИИ "ЗАЧТЕНО":

Работа получает "зачтено", если одновременно выполнены все условия:
1. Тема исследования чётко сформулирована и имеет смысл.
2. Указано направление применения, связанное с темой.
3. Есть хотя бы краткое обоснование выбора темы.
4. Присутствует описание ожидаемых результатов.
5. Объем текста составляет не менее 60 слов содержательного содержания.
6. Текст логичен, последователен и раскрывает тему.
7. Присутствуют конкретные примеры или аргументация.
8. Оформление соответствует требованиям (структура, читаемость).

---

КРИТЕРИИ "НЕ ЗАЧТЕНО":

Работа получает "не зачтено", если выполнено хотя бы одно из условий:
1. Тема отсутствует, неясна или не раскрыта.
2. Отсутствует направление применения или аргументация выбора.
3. Текст нелогичен, противоречив или не связан с заявленной темой.
4. Текст состоит преимущественно из общих фраз без конкретики или фактов.
5. Объем текста менее 50 слов.
6. Работа не демонстрирует понимания темы, цели и ожидаемых результатов.
7. Грубые ошибки в оформлении, делающие текст трудночитаемым.

---

ПРАВИЛА ОЦЕНКИ:

1. Если есть сомнение между "зачтено" и "не зачтено", выбирай "не зачтено".
2. Незначительные орфографические и стилистические ошибки не снижают оценку, если текст осмысленный.
3. Приоритет: смысл, логика, полнота раскрытия темы, аргументация.
4. Оценивай ТОЛЬКО качество содержания и оформления, НЕ источник происхождения текста.

---

ФОРМАТ ВЫВОДА:

{{
  "result": "зачтено" или "не зачтено",
  "comment": "краткий разбор содержания, логичности изложения, полноты раскрытия темы, оформления и рекомендации студенту"
}}

Отвечай на русском языке.
"""
    else:
        # Промпт для режима С AI проверкой
        base_prompt = f"""
Ты — эксперт по оценке студенческих работ.  
Оцени строго, по существу, без поблажек, но с учётом учебного уровня.  
Результат всегда бинарный: "зачтено" или "не зачтено".  
Проверка на признаки AI-генерации встроена в логику оценки.

ОЦЕНКА УВЕРЕННОСТИ ДЕТЕКТОРА AI:
— Уверенность детектора AI-генерации: {ai_confidence or 'неизвестно'}.
— ВАЖНО: если уверенность НЕ «высокая», факт наличия признаков AI не является основанием для «не зачтено». Учитывай только прочие критерии.
— Если уверенность «высокая», это допустимое основание для «не зачтено» по признакам AI (при прочих равных).

КРИТЕРИИ ОЦЕНКИ ИЗ ШАБЛОНА:
{template_text}

РАБОТА СТУДЕНТА:
{text}

---

КРИТЕРИИ "ЗАЧТЕНО":

Работа получает "зачтено", если одновременно выполнены все условия:
1. Тема исследования чётко сформулирована и имеет смысл.
2. Указано направление применения, связанное с темой.
3. Есть хотя бы краткое обоснование выбора темы.
4. Присутствует описание ожидаемых результатов.
5. Объем текста составляет не менее 60 слов содержательного содержания.
6. Текст демонстрирует признаки самостоятельности и понимания темы:
   - присутствуют личные рассуждения, конкретные примеры, естественные ошибки или логические переходы;
   - стиль текста не выглядит шаблонным, не является энциклопедическим или идеально гладким.

---

КРИТЕРИИ "НЕ ЗАЧТЕНО":

Работа получает "не зачтено", если выполнено хотя бы одно из условий:
1. Тема отсутствует, неясна.
2. Отсутствует направление применения или аргументация выбора.
3. Текст состоит преимущественно из общих фраз без конкретики или фактов.
4. Объем текста менее 50 слов.
5. Текст имеет очевидные признаки AI-генерации:
   - идеальная структура и стиль без ошибок;
   - повтор типичных фраз вроде «повышение эффективности», «улучшение качества образования»;
   - отсутствие личных высказываний, конкретных примеров или естественных ошибок;
   - монотонный, академический или энциклопедический стиль;
   - выводы и формулировки звучат как стандартные шаблоны.
6. Работа не демонстрирует понимания темы, цели и ожидаемых результатов.

---

ПРАВИЛА ОЦЕНКИ:

1. Если есть сомнение между "зачтено" и "не зачтено", выбирай "не зачтено".
2. Орфографические и стилистические ошибки не снижают оценку, если текст осмысленный и самостоятельный.
3. Приоритет: смысл, логика, самостоятельность.
4. Наличие шаблонности, клише или безличного академического стиля — основание для "не зачтено".
5. Факт AI-генерации является основанием для "не зачтено" ТОЛЬКО при уверенности детектора «высокая». В остальных случаях это не основание.

---

ФОРМАТ ВЫВОДА:

{{
  "result": "зачтено" или "не зачтено",
  "comment": "краткий разбор содержания, аргументации, структуры, признаков AI-генерации и рекомендации студенту"
}}

Отвечай на русском языке.
"""

    extra_instructions = room_prompt.strip()
    if extra_instructions:
        prompt = f"Дополнительные пожелания к проверке:\n{extra_instructions}\n\n{base_prompt}"
    else:
        prompt = base_prompt

    # Пытаемся получить валидный ответ
    response = await answer(client, text=text, prompt=prompt, models=models)
    if response:
        result = extract_binary_result(response)
        if result:
            logger.info("Получена бинарная оценка от модели.")
            return result
        else:
            logger.warning("Ответ получен, но не удалось извлечь результат.")
    else:
        logger.error("Не удалось получить ответ от модели.")

    return None


@retry(stop=stop_after_attempt(2), wait=wait_fixed(3),
       retry=retry_if_exception_type(Exception))
async def process_submission(
    json_obj: dict,
    template_text: str,
    client: GeminiClient,
    room_prompt: str = "",
    ai_check_enabled: bool = True,
) -> dict | None:
    """
    Обрабатывает одну подачу: извлекает текст из файлов,
    отправляет его модели для оценки и сохраняет результаты.
    """
    user_dir = json_obj['user']
    file_path = json_obj['file_path']
    file_type = json_obj['file_type']
    
    try:
        # Готовим директорию для результатов по каждому файлу
        results_dir = os.path.join(user_dir, 'results')
        os.makedirs(results_dir, exist_ok=True)
        base_name = os.path.basename(file_path)
        safe_name = re.sub(r'[^a-zA-Z0-9_\-\.\u0400-\u04FF]+', '_', base_name)
        per_file_result = os.path.join(results_dir, f"{safe_name}.json")

        # Если уже есть результат для конкретного файла — пропустим переобработку
        if os.path.exists(per_file_result):
            logger.info(f"✅ Уже обработан файл: {base_name}")
            async with aiofiles.open(per_file_result, 'r', encoding='utf-8') as f:
                try:
                    data = json.loads(await f.read())
                except Exception:
                    data = None
            return data
        
        # Извлечение текста в зависимости от типа файла
        if file_type == 'html':
            student_text = await extract_text_from_html(file_path)
        elif file_type in ['docx', 'doc']:
            student_text = await extract_text_from_word(file_path)
        elif file_type == 'pdf':
            student_text = await extract_text_from_pdf(file_path)
        else:
            logger.error(f"Неподдерживаемый тип файла: {file_type}")
            return None
        
        if not student_text.strip():
            logger.warning(f"Пустой текст в файле {file_path}")
            student_text = "Работа не содержит текста"
        elif len(student_text.strip()) < 70:
            logger.warning(f"Очень короткий текст в файле {file_path} ({len(student_text)} символов)")
            # Проверяем, не является ли это просто шаблоном
            if "шаблон" in student_text.lower() or "template" in student_text.lower():
                student_text = "Предоставлен только шаблон без заполнения"
            else:
                # Добавляем предупреждение о недостаточном объеме
                student_text = f"ТЕКСТ СЛИШКОМ КОРОТКИЙ ({len(student_text)} символов): {student_text}"
        
        logger.info(f"Извлечен текст из {file_path} ({len(student_text)} символов)")
        
        # Проверка на AI-генерацию только если включена AI проверка
        ai_check = None
        ai_confidence = None
        if ai_check_enabled:
            ai_check = await check_ai_generation(client, student_text)
            if ai_check:
                ai_confidence = ai_check.get('confidence')
        
        # Получение базовой оценки от AI с учетом уверенности детектора
        # Поддержка офлайн-тестирования: если DUMMY_EVAL=true — используем детерминированную оценку
        if os.environ.get('DUMMY_EVAL', '').lower() in {'1', 'true', 'yes'}:
            clean_len = len(re.sub(r'\s+', ' ', student_text).strip())
            evaluation = {
                "result": "зачтено" if clean_len >= 400 else "не зачтено",
                "comment": f"Режим DUMMY_EVAL: длина текста {clean_len} символов. Порог 400."
            }
        else:
            evaluation = await get_binary_evaluation(
                client,
                student_text,
                template_text,
                room_prompt=room_prompt,
                ai_confidence=ai_confidence,
                ai_check_enabled=ai_check_enabled,
            )
        
        if not evaluation:
            logger.error(f"Не удалось получить оценку для {user_dir}")
            evaluation = {
                "result": "не зачтено",
                "comment": "Ошибка при получении оценки от AI"
            }
        
        # Жёсткое правило: незачет за AI только при высокой уверенности (если AI проверка включена)
        if ai_check_enabled and ai_check and ai_check.get('ai_detected', False):
            if ai_check.get('confidence') == 'высокая':
                evaluation = {
                    "result": "не зачтено",
                    "comment": evaluation.get('comment', '') + "\n\nПричина: высокая уверенность детекции AI-генерации."
                }
            else:
                # Мягкое предупреждение — не влияет на итог
                evaluation['comment'] = evaluation.get('comment', '') + "\n\nЗАМЕЧАНИЕ: возможные признаки AI-генерации (уверенность не высокая)."
        
        # Сохранение результата в per-file JSON
        result_data = {
            "student": os.path.basename(user_dir),
            "file": os.path.basename(file_path),
            "date": time.strftime('%Y-%m-%d %H:%M:%S'),
            "result": evaluation['result'],
            "comment": evaluation['comment'],
            "ai_detection": ai_check if ai_check else None,
            "checked_by": "SFEDU"
        }

        result_content = json.dumps(result_data, ensure_ascii=False, indent=2)
        async with aiofiles.open(per_file_result, 'w', encoding='utf-8') as f:
            await f.write(result_content)
        logger.info(f"Результат сохранен: {per_file_result}")

        return {
            "student": os.path.basename(user_dir),
            "file": os.path.basename(file_path),
            "result": evaluation['result'],
            "comment": evaluation['comment'],
            "result_file": per_file_result
        }
        
    except Exception as e:
        logger.error(f"Ошибка при обработке подачи {user_dir}: {e}")
        return None


async def find_all_submissions(root_dir: str) -> list:
    """
    Обходит корневую директорию для поиска всех подач пользователей.
    """
    res = []
    
    for path, _, files in os.walk(root_dir):
        # Пропускаем корневую директорию
        if path == root_dir:
            continue
            
        # Ищем файлы заданий
        for file in files:
            if file.endswith(('.docx', '.doc')):
                res.append({
                    'user': path,
                    'file_path': os.path.join(path, file),
                    'file_type': 'docx' if file.endswith('.docx') else 'doc'
                })
            elif file == 'onlinetext.html':
                res.append({
                    'user': path,
                    'file_path': os.path.join(path, file),
                    'file_type': 'html'
                })
            elif file.lower().endswith('.pdf'):
                res.append({
                    'user': path,
                    'file_path': os.path.join(path, file),
                    'file_type': 'pdf'
                })
    
    logger.info(f"Всего найдено подач: {len(res)}")
    return res


@retry(stop=stop_after_attempt(2), wait=wait_fixed(3),
       retry=retry_if_exception_type(Exception))
def process_all_submissions(
    json_user_files: List[dict],
    template_text: str,
    client: GeminiClient,
    room_prompt: str = "",
    ai_check_enabled: bool = True,
    progress_callback: Callable[[str, int, int], None] | None = None,
):
    """
    Обрабатывает все подачи одновременно с ограничением на количество одновременно выполняемых задач.
    """
    semaphore = threading.Semaphore(1)  # Ограничение до 1 одновременного потока для бесплатного API
    results = []
    results_lock = threading.Lock()
    progress_lock = threading.Lock()
    total_items = len(json_user_files)
    completed = 0

    progress_bar = None
    if progress_callback is None and total_items:
        progress_bar = tqdm(total=total_items, desc="Обработка подач")
    elif progress_callback is not None:
        progress_callback("processing_submissions", 0, total_items)

    def worker(submission, template_text, client, room_prompt, ai_check_enabled):
        nonlocal results, completed
        semaphore.acquire()
        try:
            # Создаём новый цикл событий для каждого потока
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            result = loop.run_until_complete(
                process_submission(submission, template_text, client, room_prompt, ai_check_enabled),
            )
            loop.close()
            with results_lock:
                results.append(result)
        except Exception as e:
            logger.error(f"Ошибка при обработке подачи {submission.get('user')}: {e}")
            with results_lock:
                results.append(None)
        finally:
            if progress_bar is not None:
                progress_bar.update(1)
            with progress_lock:
                completed += 1
                current_completed = completed
            if progress_callback is not None:
                progress_callback("processing_submissions", current_completed, total_items)
            semaphore.release()

    threads = []

    for submission in json_user_files:
        thread = Thread(target=worker, args=(submission, template_text, client, room_prompt, ai_check_enabled))
        thread.start()
        threads.append(thread)

    # Ожидание завершения всех потоков
    for thread in threads:
        thread.join()

    if progress_bar is not None:
        progress_bar.close()
    logger.info("Все подачи были обработаны.")
    return results


async def generate_final_summary(root_dir: str):
    """
    Формирует итоговую ведомость в корневой директории.
    """
    res = []
    dedup_records = {}
    
    for path, _, files in os.walk(root_dir):
        if path == root_dir:
            continue

        # Собираем все per-file результаты
        candidates = []
        results_dir = os.path.join(path, 'results')
        if os.path.isdir(results_dir):
            for entry in os.listdir(results_dir):
                if not entry.lower().endswith('.json'):
                    continue
                result_file = os.path.join(results_dir, entry)
                try:
                    async with aiofiles.open(result_file, 'r', encoding='utf-8') as f:
                        content = await f.read()
                    data = json.loads(content)
                    student_name = data.get('student', os.path.basename(path))
                    result = data.get('result', 'не определено')
                    comment = data.get('comment', 'нет комментария')
                    date_str = data.get('date')

                    # Извлекаем полную информацию о детекции AI
                    ai_detection = data.get('ai_detection')
                    if ai_detection:
                        ai_detected = ai_detection.get('ai_detected', False)
                        ai_confidence = ai_detection.get('confidence', 'неизвестно')
                        ai_reasons = ai_detection.get('reasons', [])
                        ai_comment = ai_detection.get('comment', '')

                        if ai_detected:
                            ai_status = f"Да ({ai_confidence})"
                            if ai_comment:
                                ai_details_value = ai_comment.strip()
                            elif ai_reasons:
                                ai_details_value = f"Причины: {'; '.join(ai_reasons)}"
                            else:
                                ai_details_value = "Признаки AI-генерации зафиксированы"
                        else:
                            ai_status = "Нет"
                            if ai_comment:
                                ai_details_value = ai_comment.strip()
                            else:
                                ai_details_value = "Признаки AI-генерации не обнаружены"
                    else:
                        ai_status = "Не проверено"
                        ai_details_value = "Проверка не выполнялась"

                except json.JSONDecodeError:
                    # Fallback для старого формата
                    student_name = os.path.basename(path)
                    result_match = re.search(r'РЕЗУЛЬТАТ: (.+)', content)
                    result = result_match.group(1).strip() if result_match else "не определено"
                    comment_match = re.search(r'КОММЕНТАРИЙ:\s*(.+?)(?=\n\n|\Z)', content, re.DOTALL)
                    comment = comment_match.group(1).strip() if comment_match else "нет комментария"
                    ai_status = "Не проверено"
                    ai_details_value = "Проверка не выполнялась"
                    date_str = None

                # Служебные поля для дедупликации (не попадут в итоговый Excel)
                # 1) Попытка извлечь ID из имени студента/каталога
                raw_student = student_name or os.path.basename(path)
                folder_name = os.path.basename(path)
                student_id = None
                # Основной шаблон Moodle: ФИО_ИД_assignsubmission_<тип>
                m = re.match(r'^(.+?)_(\d+)_assignsubmission_(\w+)$', raw_student)
                if m:
                    student_id = m.group(2)
                else:
                    # Альтернативные шаблоны ID в тексте
                    m = re.search(r'ID[:=]\s*(\w{3,})', raw_student)
                    if not m:
                        m = re.search(r'\[(?:id|ID)=(\w+)\]', raw_student)
                    if m:
                        student_id = m.group(1)
                if not student_id:
                    # Попытка из имени папки
                    m = re.match(r'^(.+?)_(\d+)_assignsubmission_(\w+)$', folder_name)
                    if m:
                        student_id = m.group(2)

                # Нормализация ФИО для ключа
                try:
                    import unicodedata
                    normalized_student = unicodedata.normalize('NFC', str(raw_student).strip()).casefold()
                except Exception:
                    normalized_student = str(raw_student).strip().lower()

                # Парс даты/mtime как fallback
                from datetime import datetime
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                    except Exception:
                        parsed_date = None
                try:
                    mtime = os.path.getmtime(result_file)
                except Exception:
                    mtime = 0.0

                    # Ключ дедупликации
                    key = student_id if student_id else normalized_student

                    # Сохраняем запись с техническими полями (кандидат)
                    record = {
                        'Студент': raw_student,
                        'Результат': result,
                        'AI-детекция': ai_status,
                        'AI-детали': ai_details_value,
                        'Комментарий': comment,
                        'Путь к файлу': result_file,
                        '__student_id__': student_id,
                        '__normalized_student__': normalized_student,
                        '__parsed_date__': parsed_date,
                        '__mtime__': mtime,
                    }
                    candidates.append((key, record))
                except Exception as e:
                    logger.error(f"Ошибка при чтении {result_file}: {e}")

        # Также учитываем legacy result.txt (если он был ранее сохранён логикой до многодокового режима)
        legacy_result = os.path.join(path, 'result.txt')
        if os.path.exists(legacy_result):
            try:
                async with aiofiles.open(legacy_result, 'r', encoding='utf-8') as f:
                    content = await f.read()
                data = json.loads(content)
                raw_student = data.get('student', os.path.basename(path))
                result = data.get('result', 'не определено')
                comment = data.get('comment', 'нет комментария')
                date_str = data.get('date')
                ai_status = data.get('ai_detection', {}).get('ai_detected', 'неизвестно') if isinstance(data.get('ai_detection'), dict) else 'неизвестно'

                folder_name = os.path.basename(path)
                # Поиск ID студента
                student_id = None
                m = re.match(r'^(.+?)_(\d+)_assignsubmission_(\w+)$', raw_student)
                if m:
                    student_id = m.group(2)
                if not student_id:
                    m = re.match(r'^(.+?)_(\d+)_assignsubmission_(\w+)$', folder_name)
                    if m:
                        student_id = m.group(2)

                try:
                    import unicodedata
                    normalized_student = unicodedata.normalize('NFC', str(raw_student).strip()).casefold()
                except Exception:
                    normalized_student = str(raw_student).strip().lower()

                from datetime import datetime
                parsed_date = None
                if date_str:
                    try:
                        parsed_date = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                    except Exception:
                        parsed_date = None
                try:
                    mtime = os.path.getmtime(legacy_result)
                except Exception:
                    mtime = 0.0

                key = student_id if student_id else normalized_student
                record = {
                    'Студент': raw_student,
                    'Результат': result,
                    'AI-детекция': ai_status,
                    'AI-детали': 'legacy result.txt',
                    'Комментарий': comment,
                    'Путь к файлу': legacy_result,
                    '__student_id__': student_id,
                    '__normalized_student__': normalized_student,
                    '__parsed_date__': parsed_date,
                    '__mtime__': mtime,
                }
                candidates.append((key, record))
            except Exception as e:
                logger.error(f"Ошибка при чтении legacy result.txt: {e}")

        # Если нет кандидатов — переход к следующей папке
        if not candidates:
            continue

        # Выбор лучшей записи по ключу студента
        def record_priority(rec):
            is_pass = 1 if str(rec.get('Результат', '')).strip().lower() == 'зачтено' else 0
            dt = rec.get('__parsed_date__')
            ts = dt.timestamp() if dt else rec.get('__mtime__', 0.0)
            comment_len = len(str(rec.get('Комментарий', '')))
            return (is_pass, ts, comment_len)

        for key, record in candidates:
            best = dedup_records.get(key)
            if best is None or record_priority(record) > record_priority(best):
                dedup_records[key] = record

        # Записываем агрегированный лучший результат обратно в result.txt
        for key, record in dedup_records.items():
            if record.get('Путь к файлу', '').startswith(path):
                result_file_path = os.path.join(path, 'result.txt')
                aggregate = {
                    "student": record.get('Студент'),
                    "result": record.get('Результат'),
                    "comment": record.get('Комментарий'),
                    "date": time.strftime('%Y-%m-%d %H:%M:%S'),
                    "ai_detection": None,
                }
                try:
                    async with aiofiles.open(result_file_path, 'w', encoding='utf-8') as f:
                        await f.write(json.dumps(aggregate, ensure_ascii=False, indent=2))
                except Exception as e:
                    logger.error(f"Не удалось сохранить агрегированный result.txt: {e}")

    # Финализация: формируем список записей и оставляем только 6 колонок
    final_records = list(dedup_records.values())
    final_records.sort(key=lambda row: str(row.get('Студент', '')).casefold())
    columns = ['Студент', 'Результат', 'AI-детекция', 'AI-детали', 'Комментарий', 'Путь к файлу']
    df = pd.DataFrame([{k: rec.get(k) for k in columns} for rec in final_records], columns=columns)
    summary_path = os.path.join(root_dir, f'Итоговая_ведомость_{os.path.basename(root_dir)}.xlsx')
    df.to_excel(summary_path, index=False)
    
    # Форматирование Excel файла
    try:
        from openpyxl import load_workbook
        from openpyxl.styles import Alignment, PatternFill
        from openpyxl.utils import get_column_letter
        
        wb = load_workbook(summary_path)
        ws = wb.active
        
        # Настройка стилей
        header_fill = PatternFill(start_color="87CEEB", end_color="87CEEB", fill_type="solid")
        
        # Настройка ширины столбцов и стилей
        for col_idx, column_cells in enumerate(ws.iter_cols(min_row=1, max_row=1, min_col=1, max_col=ws.max_column), start=1):
            max_length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
            adjusted_width = max_length + 2
            column_letter = get_column_letter(col_idx)
            ws.column_dimensions[column_letter].width = adjusted_width
            for cell in column_cells:
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        # Применение переноса текста ко всем ячейкам
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True)
        
        wb.save(summary_path)
        logger.info(f"Итоговая ведомость сохранена: {summary_path}")
        
    except Exception as e:
        logger.error(f"Не удалось отформатировать Excel файл: {e}")
    
    return df, summary_path


async def check_processed_students(root_dir: str, json_user_files: list) -> tuple:
    """
    Проверяет, сколько студентов уже обработано.
    """
    processed_count = 0
    
    for submission in json_user_files:
        user_dir = submission['user']
        result_file = os.path.join(user_dir, 'result.txt')
        
        if os.path.exists(result_file):
            processed_count += 1
    
    return processed_count, 0  # Второй параметр для совместимости


async def run_auto_checker_async(
    root_dir: str,
    template_path: str,
    room_prompt: str = "",
    ai_check_enabled: bool = True,
    progress_callback: Callable[[str, int, int], None] | None = None,
) -> tuple[pd.DataFrame | None, str | None]:
    """Запускает проверку и возвращает DataFrame и путь к ведомости."""

    logger.info("🚀 Запуск проверки для директории %s", root_dir)
    logger.info("AI проверка: %s", "включена" if ai_check_enabled else "отключена")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Файл шаблона не найден: {template_path}")

    try:
        template_text = await extract_text_from_word(template_path)
        if not template_text.strip():
            template_text = "Критерии оценки не определены. Оцените работу по общим требованиям."
            logger.warning("Шаблон пустой, используем общие требования.")
    except Exception as exc:  # noqa: BLE001
        logger.error("Не удалось загрузить шаблон: %s", exc)
        template_text = "Критерии оценки не определены. Оцените работу по общим требованиям."

    json_user_files = await find_all_submissions(root_dir)
    if not json_user_files:
        logger.warning("Не найдены подачи для обработки в %s", root_dir)
        return None, None

    total_submissions = len(json_user_files)
    if progress_callback is not None:
        progress_callback("collecting_submissions", 0, total_submissions)

    processed_count, _ = await check_processed_students(root_dir, json_user_files)
    remaining_count = len(json_user_files) - processed_count

    logger.info(
        "Всего работ: %s, обработано ранее: %s, осталось: %s",
        len(json_user_files),
        processed_count,
        remaining_count,
    )

    client = GeminiClient()
    process_all_submissions(
        json_user_files,
        template_text,
        client,
        room_prompt=room_prompt,
        ai_check_enabled=ai_check_enabled,
        progress_callback=progress_callback,
    )

    if progress_callback is not None:
        progress_callback("generating_summary", 0, 1)
    df, summary_path = await generate_final_summary(root_dir)
    if progress_callback is not None:
        progress_callback("generating_summary", 1, 1)
        progress_callback("finished", total_submissions, total_submissions)
    logger.info("Итоговая ведомость сохранена: %s", summary_path)
    return df, summary_path


def run_auto_checker(
    root_dir: str,
    template_path: str,
    room_prompt: str = "",
    ai_check_enabled: bool = True,
    progress_callback: Callable[[str, int, int], None] | None = None,
) -> str | None:
    """Синхронный фасад для запуска проверки из веб-приложения."""

    _df, summary_path = asyncio.run(
        run_auto_checker_async(
            root_dir,
            template_path,
            room_prompt=room_prompt,
            ai_check_enabled=ai_check_enabled,
            progress_callback=progress_callback,
        ),
    )
    return summary_path


async def main():
    """
    Основная функция для обработки всех студентов.
    """
    print("🚀 Автоматическая проверка заданий студентов")
    print("=" * 50)
    
    # Автоматические пути
    template_path = "Шаблон.docx"
    root_dir = "Примеры"
    
    print(f"📄 Шаблон: {template_path}")
    print(f"📁 Папка студентов: {root_dir}")
    
    if not os.path.exists(template_path):
        logger.error(f"Файл шаблона не найден: {template_path}")
        print("💡 Создайте файл шаблона с критериями оценки")
        return
    
    if not os.path.exists(root_dir):
        logger.error(f"Директория не найдена: {root_dir}")
        print("💡 Убедитесь, что папка с работами студентов существует")
        return
    
    df, summary_path = await run_auto_checker_async(root_dir, template_path)
    if df is None or summary_path is None:
        logger.info("Процесс завершён без формирования ведомости.")
    else:
        logger.info("✅ Обработка завершена! Отчёт: %s", summary_path)


if __name__ == "__main__":
    print("🚀 Автоматическая проверка заданий с Gemini API (Word/HTML)")
    print(f"📊 Доступно ключей: {len(API_KEYS)}")
    print("⚡ Задержки: 10-20 секунд")
    print("🔄 Автоматическая ротация ключей при квотах")
    print("📝 Бинарная оценка: зачтено/не зачтено")
    print("📄 Результаты сохраняются в result.txt")
    print("=" * 60)
    
    asyncio.run(main())
